# -*- coding: utf-8 -*-
"""
Created on Mon Apr 17 10:12:25 2023
@author: pdelabarra
"""
import pandas as pd
import docx

## Leyendo los datos desde la ficha y convirtiendo datos
data = pd.read_excel("Ficha_380_Matriz.xlsx", sheet_name = 'KEY')
calificado = data.loc[data['KEY'] == 'CL_CALIFICADO', 'Data'].values[0]
tipo_persona = data.loc[data['KEY'] == 'TIPO_PERSONA', 'Data'].values[0]
nombres = data.loc[data['KEY'] == 'CL_NOMBRES', 'Data'].values[0]
apellidos = data.loc[data['KEY'] == 'CL_APELLIDOS', 'Data'].values[0]
data = data.fillna("")
data["Data"] = data["Data"].astype(str)

## leyendo los template posibles de contrato 380
if tipo_persona == 'PERSONA_JURIDICA':
    if calificado == 'SI':
        plantilla = "380 Calif. PJ 2023 - Calificado.docx"
        name = "380 Calif. PJ 2023 - Calificado "+str(nombres)+".docx"
    else:
        plantilla = "380 Calif. PJ 2023 - NO Calificado.docx"
        name = "380 Calif. PJ 2023 - NO Calificado "+str(nombres)+".docx"
else:
    if calificado == 'SI':
        plantilla = "380 Calif. PN 2023 - Calificado.docx"
        name = "380 Calif. PN 2023 - Calificado "+str(apellidos)+".docx"
    if calificado == 'NO':
        plantilla = "380 Calif. PN 2023 - NO Calificado.docx"
        name = "380 Calif. PN 2023 - NO Calificado "+str(apellidos)+".docx"
doc = docx.Document(plantilla)

## funcion para calcular el perfil del cliente segun encuesta
def perfil_mbi(df):
    df = df[df["Puntaje"] != ""]
    puntaje = df[df["Data"] == "X"]['Puntaje'].sum()
    p_4 = data.loc[data['KEY'] == 'EI_PREGUNTA_4_Ganar 440.000 y no perder nada',
                   'Data'].values[0]
    p_6 = data.loc[data['KEY'] == 'EI_PREGUNTA_6_No, abandonaría la inversión inmediatamente',
                   'Data'].values[0]
    p_3 = data.loc[data['KEY'] == 'EI_PREGUNTA_3_No, nunca he invertido',
                   'Data'].values[0]
    p_9 = data.loc[data['KEY'] == 'EI_PREGUNTA_9_Preservar el capital invertido',
                   'Data'].values[0]
    p_1 = data.loc[data['KEY'] == 'EI_PREGUNTA_1_Mayor a 65 años',
                   'Data'].values[0]
    if p_6 == 'X':
        return "Muy Conservador"
    if p_4 == 'X' and p_3 == 'X':
        return "Muy Conservador"
    if p_4 == 'X' and p_9 == 'X':
        return "Muy Conservador"
    if p_4 == 'X' and p_1 == 'X':
        return "Muy Conservador"
    if p_3 == 'X' and p_9 == 'X':
        return "Muy Conservador"
    else:
        if puntaje < 11:
            return "Muy Conservador"
        if 10 < puntaje < 16:
            return "Conservador"
        if 15 < puntaje < 28:
            return "Balanceado"
        if 27 < puntaje < 41:
            return "Agresivo"

## Funcion para formatear los datos en caso de ser fecha u otros
def format_value(value):
    try:
        date_obj = pd.to_datetime(value, format="%Y-%m-%d")
        return date_obj.strftime("%Y-%m-%d")
    except ValueError:
        if value == 0 or value == "0":
            return ""
        else:
            return value

## Funcion que reemplaza las variables en documentos
def replace_variable(doc, variable_dict):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for variable, value in variable_dict.items():
                if variable in run.text:
                    aux = run.text
                    words = run.text.split()
                    for i in range(len(words)):
                        if words[i] == variable:
                            # replace the word with 'A'
                            words[i] = value
                    # join the words to form the modified paragraph text
                    modified_text = ' '.join(words)
                    # run.text = run.text.replace(aux, modified_text)
                    run.text = ' '+modified_text+' '
        # if any(variable in paragraph.text for variable in variable_dict.keys()):
        #     for variable, value in variable_dict.items():
        #         if variable in paragraph.text:
        #             # for run in paragraph.runs.split(' '):
        #             paragraph.text = paragraph.text.replace(variable, format_value(value))
    for tables in doc.tables:
        for row in tables.rows:
            for cell in row.cells:
                aux = cell.text
                if aux in variable_dict.keys():
                    cell.text = cell.text.replace(aux, variable_dict[aux])  

## Crear el diccionario con la data a reemplazar
variable_dict = {}
for index, row in data.iterrows():
    variable_name = data.iloc[index]['KEY']
    valor = data.iloc[index]['Data']
    variable_dict[variable_name] = format_value(valor)
variable_dict["CL_PERFIL"] = perfil_mbi(data)

## Reemplazar las variales en el documento
replace_variable(doc, variable_dict)
doc.save(name)